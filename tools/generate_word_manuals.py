from __future__ import annotations

import os
import re
from docx import Document
from docx.shared import Pt


OUT_DIR = r"c:\xampp\htdocs\CRMnv\docs\manuales_entrega"

CLIENTE_EMPRESA = "CE consultoría y capacitación empresarial"
MODALIDAD = "Administrador de negocio"
VERSION = "v1.0"
FECHA = "28/03/2026"
ELABORO = "Zamanta Columba Silva Palacios"


cover_lines = [
    "CE consultoría y capacitación empresarial (Corporativo empresarial, Juárez & Reyes)",
    "CRMnv — Sistema de Gestión de Relación con el Cliente (CRM)",
    "",
    "MANUALES DE OPERACIÓN Y DOCUMENTACIÓN TÉCNICA",
    "- Manual 1/4: Propiedad y Control",
    "- Manual 2/4: Manual del Usuario (Operación)",
    "- Manual 3/4: Manual del Administrador (Configuración)",
    "- Manual 4/4: Documentación Técnica (El \"Código\")",
    "",
    f"Cliente / Empresa: {CLIENTE_EMPRESA}",
    f"Modalidad: {MODALIDAD}",
    f"Versión: {VERSION}",
    f"Fecha: {FECHA}",
    f"Elaboró: {ELABORO}",
    "Revisó: _______________________________",
]


manual_1 = r"""
# Manual 1/4 — Propiedad y Control

## 1. Declaración de Propiedad
El presente documento y el sistema desarrollado para el CRMnv (incluyendo su base de datos, código fuente, vistas, configuraciones, plantillas, documentos generados y componentes asociados) pertenecen legalmente a:

CE consultoría y capacitación empresarial (Corporativo empresarial, Juárez & Reyes).

La entrega se realiza bajo uso autorizado. Queda entendido que el cliente reconoce que el sistema es una obra protegida, y cualquier reproducción, distribución o puesta a disposición de terceros no autorizados puede constituir una infracción.

## 2. Acceso Maestro
Se deberán entregar al cliente las credenciales de “Súper Administrador” (Administrador con control total), con el objetivo de permitir que el cliente pueda:

1. Administrar altas y bajas de usuarios (cuando aplique).
2. Revisar, atender y resolver aprobaciones pendientes.
3. Realizar respaldo/exportación de información.
4. Gestionar mantenimiento operativo del sistema.

Buenas prácticas (recomendación):
- La entrega de credenciales debe realizarse por un medio seguro.
- Se deberá llevar registro de la fecha de entrega y responsable.

Captura recomendada (Figura M1-1):
- Captura del proceso/constancia de entrega de credenciales (sin mostrar contraseñas).
""".strip()


manual_2 = r"""
# Manual 2/4 — Manual del Usuario (Operación)

## 1. Alcance
Este manual está dirigido a usuarios del sistema (operación diaria). Describe cómo:

- Registrar prospectos en forma de Empresas y Contactos.
- Gestionar el embudo/avance mediante el estado de prospecto (semáforo).
- Registrar y cerrar ventas mediante Historial de Ventas y la Ficha de Venta.
- Usar Seguimientos (alarmas y bitácora de notas).
- Usar el Centro de Notificaciones (incluye recordatorios y acciones).

## 2. Vistas — Paso por paso (por pantalla)

### Vista 1. Inicio / Dashboard del usuario
Ubicación: user/dashboard

1. Inicia sesión.
2. Revisa los indicadores principales: Seguimientos pendientes, Alarmas programadas y Solicitudes pendientes (si aplica).
3. En “Mis Empresas”, busca por nombre y abre “Ver” cuando corresponda.
4. En “Mis Contactos”, revisa los recientes y usa “Ver todos los contactos”.
5. Para ventas, abre “Historial de Ventas”.

Captura recomendada (Figura M2-1): sección superior con métricas + paneles “Mis Empresas/Mis Contactos”.

### Vista 2. Empresas (listado y búsqueda)
Ubicación: user/companies/index

1. Ingresa al menú “Empresas”.
2. Usa “Buscar por nombre de empresa”.
3. Aplica con “Buscar” o limpia con “Limpiar”.
4. Revisa la tabla y usa “Ver” / “Editar” (si tu rol lo permite).

Captura recomendada (Figura M2-2): buscador + tabla (columna “Acciones”).

### Vista 3. Nueva Empresa (alta)
Ubicación: user/companies/create

1. Da clic en “Nueva Empresa”.
2. Llena los campos obligatorios: Nombre Comercial y Sector/Giro.
3. Completa campos opcionales (RFC, ubicación, datos fiscales) si aplica.
4. Selecciona el “Estado de prospecto” (semáforo).
5. Da clic en “Guardar (Pendiente de aprobación)”.
6. Confirma que el sistema indica que quedará “Pendiente” hasta aprobación.

Captura recomendada (Figura M2-3): formulario completo de “Nueva Empresa”.

### Vista 4. Detalle de Empresa
Ubicación: user/companies/show

1. Revisa “Información de la Empresa”.
2. Consulta secciones: Contactos, Historial de Ventas y Seguimientos.
3. En caso de baja (si aparece en tu rol), usa “Solicitar eliminación”.
4. Usa “Volver” para regresar.

Captura recomendada (Figura M2-4): bloques “Contactos / Historial de Ventas / Seguimientos”.

### Vista 5. Editar Empresa
Ubicación: user/companies/edit

1. Abre la empresa y selecciona “Editar”.
2. Ajusta campos: sector/giro, ubicación, ejecutivo asignado, datos fiscales.
3. Actualiza el “Estado de prospecto”.
4. Guarda con “Guardar cambios”.

Captura recomendada (Figura M2-5): pantalla “Editar Empresa” con el semáforo visible.

### Vista 6. Contactos (directorio)
Ubicación: user/contacts/index

1. Ingresa a “Contactos”.
2. Usa “Buscar por nombre de contacto”.
3. Limpia con “Limpiar” si aplica.
4. Revisa la tabla y usa “Ver” para detalle.

Captura recomendada (Figura M2-6): tabla de contactos + buscador.

### Vista 7. Nuevo Contacto
Ubicación: user/contacts/create

1. Da clic en “Nuevo Contacto”.
2. Selecciona la Empresa (obligatoria).
3. Llena “Nombre Completo” y “Correo electrónico”.
4. Completa campos opcionales (teléfono/celular, extensión, ubicación, género, etc.).
5. Llena “Datos para ficha de registro del cliente” (razón social, domicilio fiscal, RFC y régimen).
6. Guarda con “Guardar”.

Captura recomendada (Figura M2-7): formulario con “Empresa *” y bloque “Datos para ficha…”.

### Vista 8. Detalle de Contacto
Ubicación: user/contacts/show

1. Revisa la información del contacto y el vínculo con su empresa.
2. Gestiona “Correo” si tu rol lo permite.
3. Consulta “Notas” y edítalas si tu rol lo permite.
4. Usa “Volver” o “Editar”.

Captura recomendada (Figura M2-8): secciones “Correo” + “Notas”.

### Vista 9. Editar Contacto
Ubicación: user/contacts/edit

1. Abre el contacto y selecciona “Editar”.
2. Ajusta campos y datos para ficha.
3. Configura si se muestra el correo en fichas/listados/PDF.
4. Guarda con “Actualizar”.

Captura recomendada (Figura M2-9): formulario “Editar Contacto” con datos de ficha.

### Vista 10. Seguimientos (listado y filtros)
Ubicación: user/follow-ups/index

1. En “Filtros”, define Estado (Pendientes/Completados).
2. Selecciona el Tipo (Llamada/Reunión/Cierre) si lo requieres.
3. Aplica con “Filtrar”.
4. Revisa lista (empresa/contacto, fecha y estatus).
5. Abre detalle con “Ver” y, si aplica, completa con “Completar”.

Captura recomendada (Figura M2-10): panel filtros + lista de seguimientos.

### Vista 11. Nuevo Seguimiento
Ubicación: user/follow-ups/create

1. Da clic en “Nuevo Seguimiento”.
2. Selecciona Empresa y/o Contacto (opcional).
3. Elige “Tipo de Acción *”.
4. Define “Fecha y Hora Programada *”.
5. Define “Asignado a” si aplica.
6. Llena “Bitácora de Notas”.
7. Confirma con “Aceptar”.

Captura recomendada (Figura M2-11): formulario “Registrar nuevo seguimiento”.

### Vista 12. Detalle de Seguimiento
Ubicación: user/follow-ups/show

1. Revisa tipo, estado, fecha programada y asignación.
2. Consulta la bitácora de notas.
3. Si no está completado, utiliza “Marcar como Completado”.
4. Usa “Editar” o “Volver”.

Captura recomendada (Figura M2-12): vista detalle con botón “Marcar como Completado”.

### Vista 13. Editar Seguimiento
Ubicación: user/follow-ups/edit

1. Abre el seguimiento y selecciona “Editar”.
2. Actualiza empresa/contacto (opcional), tipo, fecha, asignado y bitácora.
3. Guarda con “Actualizar”.

Captura recomendada (Figura M2-13): formulario de edición de seguimiento.

### Vista 14. Historial de Ventas
Ubicación: user/sales/index

1. Filtra por fecha (Todos, últimos 7/14 días o último mes).
2. Filtra por empresa.
3. Limpia con “Limpiar filtros” si lo necesitas.
4. Revisa tabla y usa “Crear Ficha” / revisa el registro.

Captura recomendada (Figura M2-14): filtros + listado/acciones de ventas.

### Vista 15. Registrar Venta
Ubicación: user/sales/create

1. Completa el Paso 1: Nombre del curso/servicio, fecha de venta y empresa.
2. (Opcional) selecciona contacto.
3. Completa: monto, incluir IVA, tipo de pago, participantes y notas.
4. Da clic en “Continuar a datos de facturación”.
5. Ajusta “Datos de facturación”.
6. Guarda con “Guardar venta”.

Captura recomendada (Figura M2-15): pantalla de “Registrar Venta” (Paso 1 + datos de facturación).

### Vista 16. Editar Venta
Ubicación: user/sales/edit

1. Actualiza curso/servicio, fecha, empresa y contacto (si aplica).
2. Ajusta monto, incluir IVA, tipo pago, participantes y notas.
3. Revisa “Datos de facturación”.
4. Guarda con “Actualizar”.

Captura recomendada (Figura M2-16): sección de “Datos de facturación” en edición.

### Vista 17. Ficha de Venta (detalle + descargas)
Ubicación: user/sales/show

1. Abre la venta y revisa la ficha.
2. Descarga:
   - Descargar PDF
   - Word
3. Revisa: ficha de inscripción, datos de facturación y resumen financiero.
4. Si tu rol lo permite, elimina el registro.

Captura recomendada (Figura M2-17): header con botones “Descargar PDF” y “Word”.

### Vista 18. Filtros avanzados
Ubicación: filtros/index

1. Elige campos usando botones.
2. Abre el panel del filtro para seleccionar valores.
3. Define alcance: Empresa / Contacto / Ambos.
4. Aplica con “Aplicar filtros”.
5. Revisa resultados en tablas y usa “Ver”.

Captura recomendada (Figura M2-18): panel de filtros + resultados.

### Vista 19. Notificaciones y recordatorios (Centro)
Ubicación: notifications/index

1. Entra a “Notificaciones”.
2. En “Recordatorios”, usa “Agregar recordatorio” y marca como completado.
3. En notificaciones, abre detalle al hacer clic.
4. Acciones por notificación:
   - Marcar como leída (si no está leída)
   - Destacar (estrella ★/☆)
   - Eliminar
5. Acciones masivas: selecciona varias y usa acciones de lectura o eliminación.

Captura recomendada (Figura M2-19): bloque “Recordatorios” + lista de notificaciones.

### Vista 20. Perfil
Ubicación: profile/edit

1. Actualiza información del perfil.
2. Actualiza contraseña desde “Actualizar contraseña”.
3. Verifica confirmaciones.

Captura recomendada (Figura M2-20): “Actualizar contraseña”.

## 3. Guía de Procesos (End-to-End)

### A. Registrar prospecto (Empresa + Contactos)
1. Registra Empresa (queda “Pendiente”).
2. Registra contactos vinculados.
3. El admin aprueba para que se reflejen en el CRM.

Captura sugerida: empresa y contacto con estado “Pendiente”.

### B. Mover en el embudo
1. Abre la Empresa.
2. Edita “Estado de prospecto”.
3. Registra seguimientos conforme avances.

### C. Cerrar un trato
1. Ve a Historial de Ventas.
2. Registra la venta.
3. Descarga la ficha (PDF/Word).
4. Actualiza el semáforo según tu proceso.
""".strip()


manual_3 = r"""
# Manual 3/4 — Manual del Administrador (Configuración)

## 1. Alcance
Este manual está orientado a Administrador de negocio para:
- Gestionar aprobaciones.
- Exportar/importar información para continuidad operativa.
- Operar gestión de datos.
- Revisar seguridad y permisos por rol.
- Gestionar respaldos.
- Operar asistencia de contraseñas de usuarios.

## 2. Vistas — Paso por paso (por pantalla)

### Vista 1. Gestión de Datos
Ubicación: data-management/index

1. Ingresa a “Gestión de Datos”.
2. Exporta (si tienes rol admin):
   - Empresas (CSV)
   - Contactos (CSV)
   - Seguimientos (CSV)
3. Importa:
   - Selecciona archivo Excel (acepta .xlsx, .xls, .csv)
   - Da clic en “Importar base Excel”.
4. Revisa “Últimas empresas” y usa “Ver / Editar / Eliminar” (según permisos).

Captura recomendada (Figura M3-1): panel exportar/importar y tabla de últimas empresas.

### Vista 2. Solicitudes pendientes (tabs)
Ubicación: approvals/index

1. Abre “Solicitudes pendientes”.
2. Revisa el contador.
3. Selecciona tabs:
   - Empresas
   - Usuarios
   - Contactos
4. Para cada solicitud:
   - Aprobar
   - Denegar (cuando aplique con confirmación)

Captura recomendada (Figura M3-2): recuadro de total + pestañas.

### Vista 3. Aprobaciones de Empresas (legacy)
Ubicación: approvals/companies

1. Abre “Aprobaciones de empresas”.
2. Identifica si es “Eliminación solicitada” o “Nuevo registro solicitado”.
3. Usa:
   - Aprobar (alta o eliminación)
   - Denegar (si aplica)

Captura recomendada (Figura M3-3): lista de empresas con botones de acciones.

### Vista 4. Aprobaciones de Usuarios (legacy)
Ubicación: approvals/users

1. Abre “Aprobaciones de usuarios”.
2. Acepta o deniega solicitudes.

Captura recomendada (Figura M3-4): tarjeta/lista con “Aceptar / Denegar”.

### Vista 5. Perfil Admin (asistencia de contraseñas)
Ubicación: profile/edit (cuando es admin)

1. Busca usuario por nombre o correo.
2. Visualiza su información.
3. Restaura contraseña del usuario (puede generarse temporal segura si se deja vacío).
4. Entrega la contraseña temporal (visible una sola vez).

Captura recomendada (Figura M3-5): panel “Asistencia de contraseñas…”.

## 3. Personalización (campos/etapas/menús)

### Lo que sí se hace desde operación
- Ajuste del “Estado de prospecto” (semáforo) desde edición de Empresas/Contactos.

### Lo que no está como catálogo editable sin programar
- Las etiquetas/nombres del semáforo están definidas en constantes del código:
  - Company::PROSPECT_STATUS_LABELS
  - Contact::PROSPECT_STATUS_LABELS

Captura sugerida (Figura M3-6, opcional): pantalla donde se visualiza el semáforo actual.

## 4. Seguridad (permisos por rol)
El sistema limita lo que cada usuario ve y hace mediante combinación de:
- Roles/permisos (Spatie Permission).
- Políticas (Policies) por modelo.

Reglas reales (ejemplos):
- CompanyPolicy: admin ve todas; ejecutivo ve empresas creadas o aprobadas.
- ContactPolicy: admin ve todo; ejecutivo solo los creados por él.
- FollowUpPolicy: admin ve todo; ejecutivo los creados o asignados.
- SalePolicy: admin ve todo; ejecutivo solo las suyas (created_by).

Captura recomendada (Figura M3-7, opcional): evidencia de un usuario sin permiso sin acceso a acciones.

## 5. Respaldos (Backups)

Respaldo por exportación (recomendado desde app):
1. En “Gestión de Datos”, exporta CSV:
   - Empresas
   - Contactos
   - Seguimientos
2. Guarda en carpeta externa con fecha.

Captura recomendada (Figura M3-8): botones EXPORTAR.

Respaldo completo a nivel base de datos:
1. Exporta dump SQL desde phpMyAdmin o herramienta equivalente.
2. Incluye al menos tablas principales del CRM.
3. Guarda en almacenamiento externo y registra versión/fecha.

Captura recomendada (Figura M3-9, opcional): exportación DB.
""".strip()


manual_4 = r"""
# Manual 4/4 — Documentación Técnica (El “Código”)

## 1. Diccionario de Datos (tablas y campos)

### `companies`
- `id` (PK)
- `nombre_comercial` (único)
- `rfc` (único)
- `sector`, `municipio`, `estado`
- `ejecutivo_asignado`
- `datos_fiscales`
- `status_color` (semáforo)
- `approval_status` (`pendiente`/`aprobado`)
- auditoría: `created_by`, `approved_by`, `approved_at`
- eliminación: `deletion_pending`, `deletion_requested_by`, `deletion_requested_at`, `deletion_resolution`, `deletion_resolution_note`, `deletion_resolved_at`, `deletion_resolved_by`, `deletion_decision_user_id`

### `contacts`
- `id` (PK)
- `company_id` (FK)
- identidad: `nombre_completo`, `genero`, `puesto_de_trabajo`, `departamento`
- contacto: `celular`, `telefono`, `extension`, `email` (único), `email_activo`
- ubicación: `municipio`, `estado`
- notas: `notas`
- ficha/facturación: `razon_social`, `nombre_comercial`, `calle_numero`, `colonia_cp`, `rfc`, `regimen_fiscal`
- semáforo/aprobación: `status_color`, `approval_status`, `approved_by`, `approved_at`, `motivo_rechazo`
- eliminación: campos `deletion_*`

### `follow_ups`
- `id`
- relaciones: `company_id` (opcional), `contact_id` (opcional)
- `tipo_accion` (`llamada`, `reunión`, `cierre`)
- `fecha_alarma`
- `bitacora_notas`
- `completado`, `completado_at`
- `notification_sent_at`
- auditoría: `created_by`, `asignado_a`

### `sales`
- `id`
- `company_id`
- `contact_id` (opcional; comprador)
- `nombre_servicio`, `fecha_venta`, `monto`, `participantes`, `notas`
- `incluye_iva`
- `tipo_pago` y etiquetas calculadas
- facturación: `colonia_cp`, `regimen_fiscal`, `forma_pago`, `uso_cfdi`, `orden_compra`
- auditoría: `created_by`

### `sale_participants`
- `id`
- `sale_id` (FK)
- `nombre`, `email` (nullable)
- `orden`

### `reminders`
- `id`
- `user_id`
- `title`, `description`
- programación: `scheduled_for`, `start_at`, `end_at`, `all_day`, `repeat`, `deadline_at`
- estado: `is_done`
- marcas del sistema: `notification_sent_at`, `pre_notification_sent_at`, `last_recurring_notify_at`
- contexto: `extension`, `nombre_cliente`, `empresa`, `correo_electronico`, `numero_telefonico`, `area`, `puesto_trabajo`

### `notifications`
- `id` (uuid)
- `type`
- `data` (payload en texto)
- `read_at`

### `users`
- `id`, `name`, `email`, `password`, `profile_photo_path`
- `approval_status`, `approved_by`, `approved_at`
- método `esAdmin()` basado en roles `admin/administrador`

### Tablas de permisos (Spatie)
- `permissions`, `roles`
- pivots: `model_has_permissions`, `model_has_roles`, `role_has_permissions`

### `saved_filters`
- `id`, `user_id`
- `name`, `entity`
- `filter_logic`
- `filters` (JSON)

## 2. Código Fuente (lenguaje y generación)

Lenguaje principal:
- Backend: PHP (Laravel)
- Vistas/UI: Blade + JavaScript (Alpine.js) + Tailwind CSS

Generación de documentos:
- PDF usando `Pdf::loadView(...)` y una vista Blade reutilizada.
- Word (.doc) reutilizando la misma vista HTML y devolviendo response con `Content-Type: application/msword`.

Archivos clave:
- `app/Http/Controllers/SalesController.php`
  - `fichaPdf(Sale $sale)`
  - `fichaWord(Sale $sale)`
- `app/Http/Controllers/ContactController.php`
  - `generatePdf(Contact $contact)`
  - `generateWord(Contact $contact)`

## 3. Arquitectura (mapa + servidor)

Flujo lógico:
1. Navegador carga vistas (Blade + Tailwind + Alpine.js).
2. Laravel resuelve rutas (`routes/web.php`).
3. Se aplican policies y permisos.
4. Eloquent consulta/actualiza MySQL.
5. Para fichas PDF/Word, se renderiza vista Blade y se descarga el documento.

Datos a completar en Word:
- Servidor: [HOST/IP]
- Sistema operativo: [Windows/Linux]
- URL base: [APP_URL]
- Puertos: [80/443/etc]
- Usuario de servidor: [SSH/RDP user o panel]
- Credenciales BD:
  - DB_HOST
  - DB_DATABASE
  - DB_USERNAME
  - DB_PASSWORD
- Clave de la app: APP_KEY
""".strip()


def parse_and_add(doc: Document, text: str) -> None:
    # Conversión simple de markdown básico a Word:
    # - "# " => Título 1, "## " => Título 2, "### " => Título 3
    # - Listas numeradas "1." => List Number
    # - Listas con guion "- " => List Bullet
    num_re = re.compile(r"^\d+\.\s+")

    for raw_line in text.splitlines():
        line = raw_line.rstrip("\n")
        if not line.strip():
            continue
        stripped = line.lstrip()

        if stripped.startswith("# "):
            doc.add_paragraph(stripped[2:].strip(), style="Heading 1")
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style="Heading 2")
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 3")
        elif num_re.match(stripped):
            doc.add_paragraph(stripped, style="List Number")
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)


def add_cover(doc: Document) -> None:
    for i, t in enumerate(cover_lines):
        if t == "":
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph(t)
        if i == 0:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
        else:
            p.runs[0].font.size = Pt(11)
    doc.add_page_break()


def build_doc(filename: str, manual_text: str) -> str:
    doc = Document()
    add_cover(doc)
    parse_and_add(doc, manual_text)
    os.makedirs(OUT_DIR, exist_ok=True)
    out_path = os.path.join(OUT_DIR, filename)
    doc.save(out_path)
    return out_path


def main() -> None:
    targets = [
        ("Manual_1_Propiedad_y_Control.docx", manual_1),
        ("Manual_2_Manual_del_Usuario_Operacion.docx", manual_2),
        ("Manual_3_Manual_del_Administrador_Configuracion.docx", manual_3),
        ("Manual_4_Documentacion_Tecnica_el_Codigo.docx", manual_4),
    ]

    created = []
    for fname, txt in targets:
        out_path = build_doc(fname, txt)
        print(out_path)
        created.append(out_path)

    combined = "\n\n".join([manual_1, manual_2, manual_3, manual_4])
    combined_path = build_doc("Manuales_Completos_CRMnv.docx", combined)
    print(combined_path)
    created.append(combined_path)

    return created


if __name__ == "__main__":
    main()

